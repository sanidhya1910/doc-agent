"""Every artifact must reopen cleanly in the library that owns its format."""

from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pytest
from PIL import Image, ImageDraw

from docagent.export import build_bundle, build_manifest, writers
from docagent.state import Block, Box, DocKind, Document, Field, Page, Table


@pytest.fixture
def doc(tmp_path) -> Document:
    image = Image.new("RGB", (600, 260), "white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 20), "ACME INVOICE 2026-114", fill="black")
    draw.text((20, 60), "Widget 2 9.99", fill="black")

    page1 = Page(
        index=0, image=image, source="invoice.pdf", kind=DocKind.INVOICE,
        kind_confidence=0.91, backend="glm_ocr",
        markdown="# ACME Invoice\n\n| Item | Qty |\n| - | - |\n| Widget | 2 |\n",
    )
    page1.tables = [Table(header=["Item", "Qty"], rows=[["Widget", "2"]],
                          page_no=1, caption="Line items")]
    page1.fields = [Field("Invoice number", "2026-114", 1, 0.95)]
    page1.blocks = [
        Block("ACME INVOICE 2026-114", Box(20, 18, 300, 18)),
        Block("Widget 2 9.99", Box(20, 58, 200, 18)),
    ]

    page2 = Page(
        index=1, source="notes.docx", kind=DocKind.PROSE, kind_confidence=0.62,
        backend="docx", has_text_layer=True,
        markdown="## Notes\n\n- first point\n\n1. numbered\n\nPlain paragraph.",
    )

    document = Document(
        pages=[page1, page2], sources=["invoice.pdf", "notes.docx"],
        instruction="line items to a spreadsheet",
        summary="Invoice 2026-114 from ACME.", workdir=str(tmp_path),
    )
    return document


def test_xlsx_reopens_with_expected_sheets_and_rows(doc, tmp_path):
    import openpyxl

    path = writers.write_xlsx(doc, tmp_path)
    workbook = openpyxl.load_workbook(path)
    assert "Line items" in workbook.sheetnames
    assert "Fields" in workbook.sheetnames
    assert "Manifest" in workbook.sheetnames
    sheet = workbook["Line items"]
    assert [c.value for c in sheet[1]] == ["Item", "Qty"]
    assert [c.value for c in sheet[2]] == ["Widget", "2"]


def test_xlsx_sheet_names_are_excel_legal(tmp_path):
    import openpyxl

    document = Document(workdir=str(tmp_path))
    page = Page(index=0, markdown="x", has_text_layer=True)
    page.tables = [
        Table(header=["a"], rows=[["1"]], page_no=1, caption="bad/name:[with]*chars?"),
        Table(header=["a"], rows=[["2"]], page_no=1, caption="bad/name:[with]*chars?"),
    ]
    document.pages = [page]
    workbook = openpyxl.load_workbook(writers.write_xlsx(document, tmp_path))
    for name in workbook.sheetnames:
        assert len(name) <= 31
        assert not set(name) & set("[]:*?/\\")
    assert len(set(workbook.sheetnames)) == len(workbook.sheetnames)


def test_docx_preserves_tables_and_headings(doc, tmp_path):
    import docx

    document = docx.Document(writers.write_docx(doc, tmp_path))
    assert len(document.tables) == 1
    assert [c.text for c in document.tables[0].rows[0].cells] == ["Item", "Qty"]
    text = "\n".join(p.text for p in document.paragraphs)
    assert "first point" in text
    # Table rows became a real Word table, so the raw pipe text is gone.
    assert "| Widget |" not in text


def test_pdf_report_opens_and_has_pages(doc, tmp_path):
    import pikepdf

    with pikepdf.open(writers.write_pdf_report(doc, tmp_path)) as pdf:
        assert len(pdf.pages) >= 1


def test_searchable_pdf_text_matches_the_ocr_text(doc, tmp_path):
    import pypdfium2 as pdfium

    path = writers.write_searchable_pdf(doc, tmp_path)
    pdf = pdfium.PdfDocument(path)
    try:
        textpage = pdf[0].get_textpage()
        extracted = textpage.get_text_range() or ""
        textpage.close()
    finally:
        pdf.close()
    assert "ACME INVOICE 2026-114" in extracted
    assert "Widget" in extracted


def test_csv_written_per_table(doc, tmp_path):
    paths = writers.write_csv(doc, tmp_path)
    assert len(paths) == 1
    assert Path(paths[0]).read_text(encoding="utf-8-sig").splitlines()[0] == "Item,Qty"


def test_fields_json_contains_only_fields(doc, tmp_path):
    payload = json.loads(Path(writers.write_fields_json(doc, tmp_path)).read_text("utf-8"))
    assert [f["key"] for f in payload["fields"]] == ["Invoice number"]
    assert "pages" not in payload


def test_bundle_contains_every_artifact_plus_manifest(doc, tmp_path):
    zip_path = build_bundle(doc, tmp_path)
    with zipfile.ZipFile(zip_path) as archive:
        names = set(archive.namelist())
    assert {"document.md", "document.json", "manifest.json", "tables.xlsx",
            "summary.pdf", "searchable.pdf", "document.docx"} <= names
    assert not doc.warnings


def test_manifest_flags_the_low_confidence_page(doc):
    manifest = build_manifest(doc)
    assert manifest["needs_review"] is True
    assert any("page 2" in c for c in manifest["concerns"])
    assert manifest["min_page_confidence"] == pytest.approx(0.62)


def test_manifest_is_clean_when_everything_is_confident(doc):
    doc.pages[1].kind_confidence = 0.95
    manifest = build_manifest(doc)
    assert manifest["needs_review"] is False
    assert manifest["concerns"] == []


def test_a_failing_writer_does_not_abort_the_bundle(doc, tmp_path, monkeypatch):
    def boom(*args, **kwargs):
        raise RuntimeError("simulated failure")

    monkeypatch.setattr(writers, "write_xlsx", boom)
    zip_path = build_bundle(doc, tmp_path)
    with zipfile.ZipFile(zip_path) as archive:
        names = set(archive.namelist())
    assert "tables.xlsx" not in names
    assert "document.md" in names  # the rest still got written
    assert any("could not write xlsx" in w for w in doc.warnings)
