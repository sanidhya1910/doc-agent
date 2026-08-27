"""Artifact writers. All CPU, all deterministic, none of them cost GPU quota.

Each writer takes the :class:`~docagent.state.Document` blackboard and a
destination directory, and returns the path it wrote.
"""

from __future__ import annotations

import csv
import io
import json
import logging
import os
import re
from pathlib import Path

from ..state import Document, Table
from ..tables import _unique, strip_tables

log = logging.getLogger("docagent.export.writers")

__all__ = [
    "write_markdown",
    "write_text",
    "write_json",
    "write_fields_json",
    "write_jsonl",
    "write_csv",
    "write_xlsx",
    "write_docx",
    "write_pdf_report",
    "write_searchable_pdf",
]


def _pil_to_buffer(image) -> io.BytesIO:
    """reportlab wants a file-like PNG, not a PIL object."""
    buf = io.BytesIO()
    image.save(buf, format="PNG", optimize=True)
    buf.seek(0)
    return buf


#: Excel forbids these in sheet names and caps the length at 31.
_SHEET_BAD = re.compile(r"[\[\]:*?/\\]")


def _out(directory: str | os.PathLike, name: str) -> Path:
    path = Path(directory)
    path.mkdir(parents=True, exist_ok=True)
    return path / name


def _safe_sheet_name(raw: str, used: set[str]) -> str:
    name = _SHEET_BAD.sub("-", raw).strip() or "Sheet"
    name = name[:31]
    base = name
    n = 2
    while name.lower() in used:
        suffix = "_%d" % n
        name = base[: 31 - len(suffix)] + suffix
        n += 1
    used.add(name.lower())
    return name


# ---------------------------------------------------------------------------
# text-ish formats
# ---------------------------------------------------------------------------

def write_markdown(doc: Document, directory: str | os.PathLike, name: str = "document.md") -> str:
    parts: list[str] = []
    for page in doc.pages:
        parts.append("<!-- page %d (%s, %s) -->" % (page.page_no, page.source, page.kind.value))
        parts.append(page.text().strip())
    path = _out(directory, name)
    path.write_text("\n\n".join(p for p in parts if p), encoding="utf-8")
    return str(path)


def write_text(doc: Document, directory: str | os.PathLike, name: str = "document.txt") -> str:
    path = _out(directory, name)
    path.write_text(doc.full_text(), encoding="utf-8")
    return str(path)


def _document_payload(doc: Document) -> dict:
    return {
        "sources": doc.sources,
        "instruction": doc.instruction,
        "summary": doc.summary,
        "page_count": doc.n_pages,
        "overall_type": doc.dominant_kind().value,
        "warnings": doc.warnings,
        "pages": [
            {
                "page": p.page_no,
                "source": p.source,
                "type": p.kind.value,
                "confidence": round(p.kind_confidence, 4),
                "backend": p.backend,
                "has_text_layer": p.has_text_layer,
                "text": p.text(),
                "tables": [
                    {
                        "caption": t.caption,
                        "header": t.header,
                        "rows": t.rows,
                        "shape": list(t.shape),
                    }
                    for t in p.tables
                ],
                "fields": [
                    {"key": f.key, "value": f.value, "confidence": round(f.confidence, 4)}
                    for f in p.fields
                ],
                "warnings": p.warnings,
            }
            for p in doc.pages
        ],
    }


def write_json(doc: Document, directory: str | os.PathLike, name: str = "document.json") -> str:
    path = _out(directory, name)
    path.write_text(
        json.dumps(_document_payload(doc), indent=2, ensure_ascii=False), encoding="utf-8"
    )
    return str(path)


def write_fields_json(doc: Document, directory: str | os.PathLike,
                      name: str = "fields.json") -> str:
    """Just the key-value pairs, for callers that only want the form data."""
    payload = {
        "source": ", ".join(doc.sources),
        "type": doc.dominant_kind().value,
        "fields": [
            {
                "key": f.key,
                "value": f.value,
                "page": f.page_no,
                "confidence": round(f.confidence, 4),
            }
            for f in doc.all_fields()
        ],
    }
    path = _out(directory, name)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    return str(path)


def write_jsonl(doc: Document, directory: str | os.PathLike, name: str = "pages.jsonl") -> str:
    """One JSON object per page, for streaming into downstream pipelines."""
    path = _out(directory, name)
    with open(path, "w", encoding="utf-8") as fh:
        for page in _document_payload(doc)["pages"]:
            fh.write(json.dumps(page, ensure_ascii=False) + "\n")
    return str(path)


def write_csv(doc: Document, directory: str | os.PathLike, prefix: str = "table") -> list[str]:
    """One CSV per detected table. Returns every path written."""
    paths: list[str] = []
    for i, table in enumerate(doc.all_tables(), start=1):
        path = _out(directory, "%s_%02d_p%d.csv" % (prefix, i, table.page_no))
        with open(path, "w", encoding="utf-8-sig", newline="") as fh:
            writer = csv.writer(fh)
            if table.header:
                writer.writerow(table.header)
            writer.writerows(table.rows)
        paths.append(str(path))
    return paths


# ---------------------------------------------------------------------------
# xlsx
# ---------------------------------------------------------------------------

def write_xlsx(doc: Document, directory: str | os.PathLike, name: str = "tables.xlsx") -> str:
    """One sheet per table, plus Fields and Manifest sheets for provenance."""
    import openpyxl
    from openpyxl.styles import Font
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    used: set[str] = set()
    bold = Font(bold=True)

    for i, table in enumerate(doc.all_tables(), start=1):
        title = table.caption or ("Table %d (p%d)" % (i, table.page_no))
        ws = wb.create_sheet(_safe_sheet_name(title, used))
        if table.header:
            ws.append(_unique(table.header))
            for cell in ws[1]:
                cell.font = bold
            ws.freeze_panes = "A2"
        for row in table.rows:
            ws.append(list(row))
        for col in range(1, (len(table.header) or table.shape[1] or 1) + 1):
            width = 12
            for cell in ws[get_column_letter(col)]:
                if cell.value is not None:
                    width = max(width, min(len(str(cell.value)) + 2, 60))
            ws.column_dimensions[get_column_letter(col)].width = width

    fields = doc.all_fields()
    if fields:
        ws = wb.create_sheet(_safe_sheet_name("Fields", used))
        ws.append(["Field", "Value", "Page", "Confidence"])
        for cell in ws[1]:
            cell.font = bold
        ws.freeze_panes = "A2"
        for f in fields:
            ws.append([f.key, f.value, f.page_no, round(f.confidence, 3)])
        ws.column_dimensions["A"].width = 32
        ws.column_dimensions["B"].width = 52

    ws = wb.create_sheet(_safe_sheet_name("Manifest", used))
    ws.append(["Page", "Source", "Type", "Confidence", "Backend", "Text layer", "Warnings"])
    for cell in ws[1]:
        cell.font = bold
    ws.freeze_panes = "A2"
    for p in doc.pages:
        ws.append([
            p.page_no, p.source, p.kind.value, round(p.kind_confidence, 3),
            p.backend or "-", "yes" if p.has_text_layer else "no", "; ".join(p.warnings),
        ])
    for col, width in zip("ABCDEFG", (8, 28, 16, 12, 18, 12, 40)):
        ws.column_dimensions[col].width = width

    path = _out(directory, name)
    wb.save(path)
    return str(path)


# ---------------------------------------------------------------------------
# docx
# ---------------------------------------------------------------------------

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+[.)]\s+(.*)$")


def _add_table(document, table: Table) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    n_cols = len(table.header) or table.shape[1]
    if n_cols <= 0:
        return
    if table.caption:
        caption = document.add_paragraph(table.caption)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True

    doc_table = document.add_table(rows=0, cols=n_cols)
    doc_table.style = "Table Grid"
    if table.header:
        cells = doc_table.add_row().cells
        for i, text in enumerate(table.header[:n_cols]):
            cells[i].text = text
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
    for row in table.rows:
        cells = doc_table.add_row().cells
        for i, text in enumerate(list(row)[:n_cols]):
            cells[i].text = text


def write_docx(doc: Document, directory: str | os.PathLike, name: str = "document.docx") -> str:
    """Markdown structure preserved: headings, lists and tables become real Word objects."""
    import docx

    document = docx.Document()
    document.add_heading(doc.sources[0] if doc.sources else "Extracted document", level=0)
    if doc.summary:
        document.add_heading("Summary", level=1)
        document.add_paragraph(doc.summary)

    for page in doc.pages:
        document.add_heading("Page %d" % page.page_no, level=1)
        # Tables are emitted as real Word tables from the parsed structure, so
        # the raw pipe rows are stripped out of the prose to avoid duplication.
        for line in strip_tables(page.text()).splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            heading = _HEADING.match(stripped)
            if heading:
                document.add_heading(heading.group(2), level=min(len(heading.group(1)) + 1, 9))
                continue
            bullet = _BULLET.match(stripped)
            if bullet:
                document.add_paragraph(bullet.group(1), style="List Bullet")
                continue
            numbered = _NUMBERED.match(stripped)
            if numbered:
                document.add_paragraph(numbered.group(1), style="List Number")
                continue
            document.add_paragraph(stripped)

        for table in page.tables:
            _add_table(document, table)

    path = _out(directory, name)
    document.save(path)
    return str(path)


# ---------------------------------------------------------------------------
# pdf
# ---------------------------------------------------------------------------

def write_pdf_report(doc: Document, directory: str | os.PathLike, name: str = "summary.pdf") -> str:
    """A readable report: summary, key fields, tables, per-page provenance."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Image as PdfImage, PageBreak, Paragraph, SimpleDocTemplate, Spacer,
        Table as PdfTable, TableStyle,
    )
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "body", parent=styles["BodyText"], alignment=TA_LEFT, leading=14, spaceAfter=6
    )
    cell = ParagraphStyle("cell", parent=body, fontSize=8, leading=10, spaceAfter=0)

    path = _out(directory, name)
    pdf = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm,
        title="Document summary",
    )

    grid = TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B0B7C3")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF1F6")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])

    story: list = [
        Paragraph("Document summary", styles["Title"]),
        Paragraph(
            escape("Source: %s  |  %d page(s)  |  type: %s"
                   % (", ".join(doc.sources) or "upload", doc.n_pages,
                      doc.dominant_kind().value)),
            body,
        ),
        Spacer(1, 6 * mm),
    ]

    if doc.instruction:
        story += [Paragraph("Request", styles["Heading2"]),
                  Paragraph(escape(doc.instruction), body), Spacer(1, 4 * mm)]

    if doc.summary:
        story.append(Paragraph("Summary", styles["Heading2"]))
        for para in doc.summary.split("\n\n"):
            if para.strip():
                story.append(Paragraph(escape(para.strip()), body))
        story.append(Spacer(1, 4 * mm))

    fields = doc.all_fields()
    if fields:
        story.append(Paragraph("Key fields", styles["Heading2"]))
        rows = [[Paragraph("<b>Field</b>", cell), Paragraph("<b>Value</b>", cell),
                 Paragraph("<b>Page</b>", cell)]]
        rows += [
            [Paragraph(escape(f.key), cell), Paragraph(escape(f.value), cell),
             Paragraph(str(f.page_no), cell)]
            for f in fields
        ]
        table = PdfTable(rows, colWidths=[45 * mm, 105 * mm, 14 * mm], repeatRows=1)
        table.setStyle(grid)
        story += [table, Spacer(1, 4 * mm)]

    for i, tbl in enumerate(doc.all_tables(), start=1):
        story.append(Paragraph(
            escape(tbl.caption or "Table %d (page %d)" % (i, tbl.page_no)), styles["Heading2"]
        ))
        header = tbl.header or ["col %d" % (c + 1) for c in range(tbl.shape[1])]
        rows = [[Paragraph("<b>%s</b>" % escape(h), cell) for h in header]]
        # Long tables get truncated: the full data is always in the XLSX/CSV.
        for row in tbl.rows[:60]:
            padded = list(row) + [""] * (len(header) - len(row))
            rows.append([Paragraph(escape(c), cell) for c in padded[: len(header)]])
        available = 164 * mm
        table = PdfTable(rows, colWidths=[available / max(len(header), 1)] * len(header),
                         repeatRows=1)
        table.setStyle(grid)
        story.append(table)
        if len(tbl.rows) > 60:
            story.append(Paragraph(
                escape("... %d more rows; see the spreadsheet export."
                       % (len(tbl.rows) - 60)), cell))
        story.append(Spacer(1, 4 * mm))

    story.append(PageBreak())
    story.append(Paragraph("Provenance", styles["Heading2"]))
    rows = [[Paragraph("<b>%s</b>" % h, cell)
             for h in ("Page", "Source", "Type", "Conf.", "Backend", "Notes")]]
    for p in doc.pages:
        rows.append([
            Paragraph(str(p.page_no), cell),
            Paragraph(escape(p.source), cell),
            Paragraph(p.kind.value, cell),
            Paragraph("%.2f" % p.kind_confidence, cell),
            Paragraph(escape(p.backend or "-"), cell),
            Paragraph(escape("; ".join(p.warnings)), cell),
        ])
    table = PdfTable(rows, colWidths=[12 * mm, 42 * mm, 24 * mm, 14 * mm, 30 * mm, 42 * mm],
                     repeatRows=1)
    table.setStyle(grid)
    story.append(table)

    if doc.warnings:
        story += [Spacer(1, 4 * mm), Paragraph("Warnings", styles["Heading2"])]
        for w in doc.warnings:
            story.append(Paragraph(escape("- " + w), body))

    thumbs = [(p, p.thumbnail()) for p in doc.pages]
    thumbs = [(p, t) for p, t in thumbs if t is not None]
    if thumbs:
        story += [PageBreak(), Paragraph("Pages", styles["Heading2"])]
        for page, thumb in thumbs:
            width_px, height_px = thumb.size
            display_w = min(70 * mm, 164 * mm)
            display_h = display_w * height_px / max(width_px, 1)
            story.append(Paragraph(
                escape("Page %d - %s" % (page.page_no, page.source)), cell))
            story.append(PdfImage(_pil_to_buffer(thumb), width=display_w, height=display_h))
            story.append(Spacer(1, 3 * mm))

    pdf.build(story)
    return str(path)


def write_searchable_pdf(
    doc: Document, directory: str | os.PathLike, name: str = "searchable.pdf"
) -> str:
    """Original page images with an invisible, selectable text layer on top.

    Where a backend gave real line boxes (TrOCR, and any PDF text layer) the
    text is positioned over the words it came from.  GLM-OCR runs
    coordinate-free, so those pages get text distributed evenly down the page:
    the result is fully searchable and copy-pasteable, but the character
    positions are approximate rather than exact.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas as pdfcanvas

    path = _out(directory, name)
    c = pdfcanvas.Canvas(str(path))
    wrote_any = False

    for page in doc.pages:
        image = page.image
        if image is None:
            # Text-only source (Office, plain text): lay the text on a page of
            # its own so the artifact still covers every page.
            width, height = A4
            c.setPageSize((width, height))
            c.setFont("Helvetica", 9)
            y = height - 40
            for line in page.text().splitlines():
                if y < 40:
                    c.showPage()
                    c.setFont("Helvetica", 9)
                    y = height - 40
                c.drawString(40, y, line[:110])
                y -= 12
            c.showPage()
            wrote_any = True
            continue

        img_w, img_h = image.size
        c.setPageSize((img_w, img_h))
        c.drawImage(ImageReader(image.convert("RGB")), 0, 0, width=img_w, height=img_h)

        text_obj = c.beginText()
        text_obj.setTextRenderMode(3)  # invisible: selectable but not drawn

        positioned = [b for b in page.blocks if b.box is not None and b.text.strip()]
        if positioned:
            for block in positioned:
                box = block.box
                size = max(4.0, box.h * 0.8)
                text_obj.setFont("Helvetica", size)
                # PDF origin is bottom-left; image coordinates are top-left.
                text_obj.setTextOrigin(box.x, img_h - box.y - box.h + size * 0.15)
                text_obj.textLine(block.text)
        else:
            lines = [ln for ln in page.text().splitlines() if ln.strip()]
            if lines:
                step = img_h / (len(lines) + 1)
                size = max(6.0, min(step * 0.7, 18.0))
                text_obj.setFont("Helvetica", size)
                for i, line in enumerate(lines, start=1):
                    text_obj.setTextOrigin(10, img_h - step * i)
                    text_obj.textLine(line)
        c.drawText(text_obj)
        c.showPage()
        wrote_any = True

    if not wrote_any:
        c.setPageSize(A4)
        c.showPage()
    c.save()
    return str(path)
